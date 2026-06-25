from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
FORMAT_CLI = PROJECT_ROOT / "scripts" / "format_docx.py"


def run_format_cli(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(FORMAT_CLI), *args],
        cwd=PROJECT_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )


def make_docx(path: Path, lines: list[str]) -> Path:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(path))
    return path


def footer_xml(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(section.footer._element.xml for section in doc.sections)


def footer_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for section in doc.sections for paragraph in section.footer.paragraphs)


def document_xml(path: Path) -> str:
    doc = Document(str(path))
    return doc._element.xml


def document_body_block_texts(path: Path) -> list[str]:
    doc = Document(str(path))
    body = doc.element.body
    paragraph_index = 0
    table_index = 0
    blocks: list[str] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = doc.paragraphs[paragraph_index].text.strip()
            paragraph_index += 1
            if text:
                blocks.append(text)
        elif child.tag == qn("w:tbl"):
            table = doc.tables[table_index]
            table_index += 1
            first_row = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
            blocks.append(f"TABLE: {first_row}")
    return blocks


def table_grid_widths(path: Path) -> list[int]:
    doc = Document(str(path))
    grid = doc.tables[0]._tbl.tblGrid
    return [int(column.get(qn("w:w"))) for column in grid.gridCol_lst]


def test_confident_notice_fixture_generates_docx_and_report(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "一、总体要求",
            "请按要求开展安全生产检查。",
            "（一）检查范围",
            "各部门自查。",
            "测试单位",
            "2026年6月19日",
        ],
    )
    output_path = tmp_path / "notice.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    assert report_path.exists()
    assert "saved=" in result.stdout
    assert "report=" in result.stdout

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["input"] == str(input_path)
    assert report["output"] == str(output_path)
    assert report["profile_id"] == "standard-party-government"
    assert report["doc_type"] == "通知"
    assert report["structure"]["title_indices"]
    assert report["operations"]


def test_ambiguous_fixture_returns_2_and_does_not_generate_docx(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "ambiguous_source.docx",
        [
            "关于数据治理有关事项的材料",
            "现将有关情况报告如下。",
            "请贵单位协助支持。",
        ],
    )
    output_path = tmp_path / "ambiguous.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 2
    assert not output_path.exists()
    assert not report_path.exists()
    assert "你希望怎么处理" in result.stdout
    assert "通用正式文本" in result.stdout


def test_page_numbers_flag_inserts_word_page_field_and_reports_action(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "notice_source.docx",
        [
            "关于开展安全生产检查的通知",
            "各部门：",
            "请按要求开展安全生产检查。",
            "某某办公室",
            "2026年6月21日",
        ],
    )
    output_path = tmp_path / "notice_with_page_number.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--page-numbers")

    assert result.returncode == 0, result.stderr
    xml = footer_xml(output_path)
    assert "PAGE" in xml
    assert "fldChar" in xml

    report = json.loads(report_path.read_text(encoding="utf-8"))
    page_ops = [operation for operation in report["operations"] if operation["kind"] == "page_number"]
    assert page_ops
    assert page_ops[0]["params"]["existing_page_number"] is False
    assert page_ops[0]["params"]["action"] == "inserted"


def test_page_numbers_flag_preserves_non_page_footer_and_skips_insertion(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_footer.docx"
    source = Document()
    source.add_paragraph("关于开展安全检查的通知")
    source.add_paragraph("各部门：")
    source.add_paragraph("请按要求开展安全检查。")
    source.add_paragraph("某某办公室")
    source.add_paragraph("2026年6月21日")
    source.sections[0].footer.paragraphs[0].text = "内部资料"
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_footer_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--page-numbers")

    assert result.returncode == 0, result.stderr
    assert "内部资料" in footer_text(output_path)
    assert "PAGE" not in footer_xml(output_path)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    page_ops = [operation for operation in report["operations"] if operation["kind"] == "page_number"]
    assert page_ops
    assert page_ops[0]["params"]["existing_non_page_footer"] is True
    assert page_ops[0]["params"]["action"] == "skipped_non_page_footer"


def test_format_tables_flag_preserves_table_and_formats_cell_text(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_table.docx"
    source = Document()
    source.add_paragraph("关于报送整改台账的通知")
    source.add_paragraph("各部门：")
    source.add_paragraph("请按表格要求报送整改台账。")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "事项"
    table.cell(0, 1).text = "责任部门"
    table.cell(1, 0).text = "安全检查"
    table.cell(1, 1).text = "综合部"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = None
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_table_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-tables")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    assert len(output.tables) == 1
    first_run = output.tables[0].cell(0, 0).paragraphs[0].runs[0]
    assert first_run.text == "事项"
    assert first_run.font.name == "仿宋_GB2312"
    assert first_run.font.size.pt == 16

    report = json.loads(report_path.read_text(encoding="utf-8"))
    table_ops = [operation for operation in report["operations"] if operation["kind"] == "table_format"]
    assert table_ops
    assert table_ops[0]["params"]["table_count"] == 1
    assert table_ops[0]["params"]["action"] == "formatted_inline"


def test_formatting_does_not_treat_numbered_colon_heading_as_recipient_or_drop_tables(tmp_path: Path) -> None:
    input_path = tmp_path / "management_rule_source.docx"
    source = Document()
    for text in [
        "超储物资内部调剂消耗指引",
        "一、超储物资定义与识别",
        "1.物资定义：超储物资是指库龄通常在一年以内，因项目结余、采购计划变更或储备定额调整而形成的，库存数量超出实际需求但物理性能完好、具备完整使用价值的备品配件及材料。",
        "2.识别与上架标准：基层企业应通过 ERP 系统或物资管理模块，定期筛选“无动态、高库存”物资，将其列入超储物资调剂清单进行专项管理。",
        "二、组织职责与权限分工",
        "1.集团供应链管理部：作为归口管理部门，负责制定管理制度，建立信息平台，并对各单位调剂完成情况进行绩效考核。",
        "2.集团物资供应中心（调剂中心）：负责跨二级单位调剂的交易撮合、物流协调及平台调剂模块的日常运营。",
        "3.二级单位（管理中心）：负责本单位内部超储物资的认定、价格审批及所属基层企业的利库监督。",
        "4.基层企业（执行主体）：负责实物信息的实时发布、质量维护、实物交接及账务处理。",
        "三、信息发布规范",
        "1.信息发布规范：上架信息须包含物资编码、参数规格、交易价格、实物照片及技术说明书等关键要素。",
        "2.动态更新机制：当物资被锁定、消耗或因其他原因发生状态变化时，发布单位应在24小时内完成信息同步。",
        "四、“应选未选”理由审核标准",
        "1.强制利库规则：需求单位在发起采购申请前，系统将自动匹配集团范围内超储物资库。若存在同类物资，原则上必须优先调剂使用。",
        "2.合理拒选理由：",
        "（一）差异化定价逻辑",
    ]:
        source.add_paragraph(text)
    table = source.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "交易主体场景"
    table.cell(0, 1).text = "价值门槛"
    table.cell(0, 2).text = "定价基础"
    table.cell(0, 3).text = "评估要求"
    table.cell(1, 0).text = "跨单位"
    table.cell(1, 1).text = "一般物资"
    table.cell(1, 2).text = "评估价"
    table.cell(1, 3).text = "按规则执行"
    source.add_paragraph("（二）核心概念定义与流程")
    source.save(str(input_path))
    output_path = tmp_path / "management_rule.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--assume-detected-type")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts[:15] == [
        "超储物资内部调剂消耗指引",
        "一、超储物资定义与识别",
        "1.物资定义：超储物资是指库龄通常在一年以内，因项目结余、采购计划变更或储备定额调整而形成的，库存数量超出实际需求但物理性能完好、具备完整使用价值的备品配件及材料。",
        "2.识别与上架标准：基层企业应通过 ERP 系统或物资管理模块，定期筛选“无动态、高库存”物资，将其列入超储物资调剂清单进行专项管理。",
        "二、组织职责与权限分工",
        "1.集团供应链管理部：作为归口管理部门，负责制定管理制度，建立信息平台，并对各单位调剂完成情况进行绩效考核。",
        "2.集团物资供应中心（调剂中心）：负责跨二级单位调剂的交易撮合、物流协调及平台调剂模块的日常运营。",
        "3.二级单位（管理中心）：负责本单位内部超储物资的认定、价格审批及所属基层企业的利库监督。",
        "4.基层企业（执行主体）：负责实物信息的实时发布、质量维护、实物交接及账务处理。",
        "三、信息发布规范",
        "1.信息发布规范：上架信息须包含物资编码、参数规格、交易价格、实物照片及技术说明书等关键要素。",
        "2.动态更新机制：当物资被锁定、消耗或因其他原因发生状态变化时，发布单位应在24小时内完成信息同步。",
        "四、“应选未选”理由审核标准",
        "1.强制利库规则：需求单位在发起采购申请前，系统将自动匹配集团范围内超储物资库。若存在同类物资，原则上必须优先调剂使用。",
        "2.合理拒选理由：",
    ]
    assert len(output.tables) == 1
    blocks = document_body_block_texts(output_path)
    assert blocks.index("（一）差异化定价逻辑") < blocks.index("TABLE: 交易主体场景 | 价值门槛 | 定价基础 | 评估要求")
    assert blocks.index("TABLE: 交易主体场景 | 价值门槛 | 定价基础 | 评估要求") < blocks.index("（二）核心概念定义与流程")


def test_formatting_preserves_multiline_title_style(tmp_path: Path) -> None:
    input_path = make_docx(
        tmp_path / "multiline_title_source.docx",
        [
            "闲置物资调剂完善后的流程与",
            "原流程对比说明",
            "一、总体情况",
            "本材料说明流程差异。",
        ],
    )
    output_path = tmp_path / "multiline_title.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert non_empty[0].text.strip() == "闲置物资调剂完善后的流程与"
    assert non_empty[1].text.strip() == "原流程对比说明"
    for paragraph in non_empty[:2]:
        assert paragraph.alignment == 1
        assert paragraph.runs[0].font.name == "方正小标宋简体"
        assert paragraph.runs[0].font.size.pt == 22


def test_multiline_title_detection_does_not_swallow_front_matter_or_tables(tmp_path: Path) -> None:
    input_path = tmp_path / "prd_source.docx"
    source = Document()
    source.add_paragraph("供应链新增应急借料业务")
    source.add_paragraph("产品需求规格说明书")
    source.add_paragraph("版本记录")
    for index in range(4):
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = f"字段{index + 1}"
        table.cell(0, 1).text = "说明"
    source.save(str(input_path))
    output_path = tmp_path / "prd.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert non_empty[0].text.strip() == "供应链新增应急借料业务"
    assert non_empty[1].text.strip() == "产品需求规格说明书"
    assert non_empty[2].text.strip() == "版本记录"
    assert len(output.tables) == 4


def test_standard_text_document_uses_standard_branch_and_specific_layout(tmp_path: Path) -> None:
    input_path = tmp_path / "standard_text_source.docx"
    source = Document()
    for text in [
        "中华人民共和国",
        "电力企业团体标准配套稿",
        "供应链服务标准",
        "（征求意见稿）",
        "目次",
        "前    言",
        "1  范围",
        "3.1  闲置物资",
        "前    言",
        "本标准依据相关规则起草。",
        "1  范围",
        "本文件规定了供应链服务范围。",
        "3.1  闲置物资",
        "长期未使用但仍具备使用价值的物资。",
    ]:
        source.add_paragraph(text)
    table = source.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "项目"
    table.cell(0, 2).text = "要求"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "信息发布"
    table.cell(1, 2).text = "内容完整"
    source.save(str(input_path))
    output_path = tmp_path / "standard_text.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    assert texts[0] == "中华人民共和国电力企业团体标准配套稿"
    body_preface_index = max(index for index, text in enumerate(texts) if text == "前    言")
    assert texts[body_preface_index + 1].startswith("本标准依据")

    toc_chapter = next(paragraph for paragraph in non_empty if paragraph.text.strip() == "1  范围")
    toc_clause = next(paragraph for paragraph in non_empty if paragraph.text.strip() == "3.1  闲置物资")
    assert toc_chapter.runs[0].bold is True
    assert toc_clause.runs[0].bold is False
    assert table_grid_widths(output_path) == [1061, 3892, 3892]

    report = json.loads(report_path.read_text(encoding="utf-8"))
    standard_ops = [operation for operation in report["operations"] if operation["kind"] == "standard_text_format"]
    assert standard_ops
    assert report["doc_type"] == "标准规范文本"
    assert standard_ops[0]["params"]["merged_cover_label"] is True


def test_generic_formal_text_doc_type_formats_ambiguous_material_without_prompt(tmp_path: Path) -> None:
    input_path = tmp_path / "generic_material.docx"
    source = Document()
    for text in [
        "专项工作交流材料",
        "相关单位：",
        "一、基本情况",
        "有关工作正在稳步推进。",
        "（一）推进情况",
        "各项任务按照计划开展。",
        "1. 重点事项",
        "重点事项已形成阶段性成果。",
        "某某部门",
        "2026年6月23日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "generic_material_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--doc-type", "通用正式文本")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    assert "你希望" not in result.stdout

    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts[0] == "专项工作交流材料"
    assert texts[1] == "相关单位："
    assert texts[-2:] == ["某某部门", "2026年6月23日"]
    assert output.paragraphs[0].alignment == 1

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_generic_formal_text_flag_formats_without_prompt(tmp_path: Path) -> None:
    input_path = tmp_path / "generic_material_flag.docx"
    source = Document()
    source.add_paragraph("专项工作交流材料")
    source.add_paragraph("一、基本情况")
    source.add_paragraph("有关工作正在推进。")
    source.save(str(input_path))

    output_path = tmp_path / "generic_material_flag_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    assert "你希望" not in result.stdout
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_glued_single_paragraph_report_is_split_into_reasonable_blocks(tmp_path: Path) -> None:
    input_path = tmp_path / "glued_report_source.docx"
    source = Document()
    paragraph = source.add_paragraph(
        "闲置及超储专区上架规范自查报告"
        "为提高闲置专区及超储专区商品信息的规范性和专业性，进一步提升专区上架管理水平，组织开展了自查工作。"
        "现将自查情况报告如下。"
        "存在的问题"
        "经全面排查，闲置及超储专区在架物资的计量单位字段存在大量英文缩写，影响商品信息的统一展示。"
        "解决措施"
        "针对英文计量单位问题，计划采取分步整改的方式，将英文缩写统一替换为中文表述。"
        "时间计划"
        "按期完成英文单位批量替换申请的提交工作，并完成相关整改复核验收。"
        "运营专班"
    )
    paragraph.alignment = 1
    source.save(str(input_path))

    output_path = tmp_path / "glued_report_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
    assert texts == [
        "闲置及超储专区上架规范自查报告",
        "为提高闲置专区及超储专区商品信息的规范性和专业性，进一步提升专区上架管理水平，组织开展了自查工作。现将自查情况报告如下。",
        "存在的问题",
        "经全面排查，闲置及超储专区在架物资的计量单位字段存在大量英文缩写，影响商品信息的统一展示。",
        "解决措施",
        "针对英文计量单位问题，计划采取分步整改的方式，将英文缩写统一替换为中文表述。",
        "时间计划",
        "按期完成英文单位批量替换申请的提交工作，并完成相关整改复核验收。",
        "运营专班",
    ]
    assert output.paragraphs[0].alignment == 1
    assert output.paragraphs[1].alignment != 1
    for heading_text in ["存在的问题", "解决措施", "时间计划"]:
        heading = next(paragraph for paragraph in output.paragraphs if paragraph.text.strip() == heading_text)
        assert heading.paragraph_format.first_line_indent.pt == 32
        assert heading.runs[0].font.name == "黑体"
        assert heading.runs[0].font.bold is True

    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "报告"
    assert report["structure"]["title_indices"] == [0]


def test_unnumbered_headings_and_subtitle_are_structured(tmp_path: Path) -> None:
    input_path = tmp_path / "unnumbered_headings_source.docx"
    input_path = make_docx(
        input_path,
        [
            "闲置及超储专区上架规范自查报告",
            "整改情况说明",
            "为提高专区上架管理水平，组织开展了自查工作。",
            "存在的问题",
            "部分物资缺少商品图片，影响物资的正常展示和采购决策效率。",
            "解决措施",
            "已安排专人对接无图片物资的核查工作，并反馈至对应单位补充图片。",
            "时间计划",
            "按期完成无图片物资的清单排查、单位对接及图片补充系统更新。",
            "运营专班",
        ],
    )
    output_path = tmp_path / "unnumbered_headings_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    assert [paragraph.text.strip() for paragraph in non_empty[:2]] == ["闲置及超储专区上架规范自查报告", "整改情况说明"]
    assert non_empty[0].alignment == 1
    assert non_empty[1].alignment == 1
    assert non_empty[1].runs[0].font.name == "方正小标宋简体"
    for heading_text in ["存在的问题", "解决措施", "时间计划"]:
        heading = next(paragraph for paragraph in non_empty if paragraph.text.strip() == heading_text)
        assert heading.paragraph_format.first_line_indent.pt == 32
        assert heading.runs[0].font.name == "黑体"
        assert heading.runs[0].font.bold is True
    signature = non_empty[-1]
    assert signature.text.strip() == "运营专班"
    assert signature.runs[0].font.name == "仿宋_GB2312"
    assert signature.runs[0].font.bold is False


def test_spaced_single_paragraph_report_recovers_nested_hierarchy(tmp_path: Path) -> None:
    input_path = tmp_path / "spaced_glued_report_source.docx"
    source = Document()
    source.add_paragraph(
        "闲置及超储专区上架规范自查报告 "
        "为提高闲置专区及超储专区商品信息的规范性和专业性，进一步提升专区上架管理水平，根据集团商品信息管理相关要求，组织开展了闲置及超储专区上架规范自查工作。现将自查情况报告如下。 "
        "存在的问题 "
        "计量单位不符合商城商品规范 "
        "经全面排查，闲置及超储专区在架物资的计量单位字段存在大量英文缩写，各类英文单位不符合商城商品信息规范要求，影响商城信息标准化管理及商品信息的统一展示。 "
        "上架价格存在异常 "
        "在自查过程中发现，部分物资的上架价格与框架物资价格对比差异较大。 "
        "部分物资缺少商品图片 "
        "商品图片是采购人员对待上架物资规格型号进行直观判断的重要依据。 "
        "解决措施 "
        "计量单位不规范问题整改方案 "
        "针对英文计量单位问题，计划采取分步整改的方式。 "
        "上架价格异常问题整改方案 "
        "针对已标记的疑似填错记录，物资管理员将逐一联系对应单位进行价格核实。 "
        "商品图片缺失问题整改方案 "
        "已安排专人对接无图片物资的核查工作。 "
        "时间计划 "
        "按期完成英文单位批量替换申请的提交工作。 "
        "后续将定期跟踪整改进度，确保各项整改措施按计划推进落实。 "
        "运营专班"
    )
    source.save(str(input_path))
    output_path = tmp_path / "spaced_glued_report_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    assert texts == [
        "闲置及超储专区上架规范自查报告",
        "为提高闲置专区及超储专区商品信息的规范性和专业性，进一步提升专区上架管理水平，根据集团商品信息管理相关要求，组织开展了闲置及超储专区上架规范自查工作。现将自查情况报告如下。",
        "一、存在的问题",
        "（一）计量单位不符合商城商品规范",
        "经全面排查，闲置及超储专区在架物资的计量单位字段存在大量英文缩写，各类英文单位不符合商城商品信息规范要求，影响商城信息标准化管理及商品信息的统一展示。",
        "（二）上架价格存在异常",
        "在自查过程中发现，部分物资的上架价格与框架物资价格对比差异较大。",
        "（三）部分物资缺少商品图片",
        "商品图片是采购人员对待上架物资规格型号进行直观判断的重要依据。",
        "二、解决措施",
        "（一）计量单位不规范问题整改方案",
        "针对英文计量单位问题，计划采取分步整改的方式。",
        "（二）上架价格异常问题整改方案",
        "针对已标记的疑似填错记录，物资管理员将逐一联系对应单位进行价格核实。",
        "（三）商品图片缺失问题整改方案",
        "已安排专人对接无图片物资的核查工作。",
        "三、时间计划",
        "按期完成英文单位批量替换申请的提交工作。",
        "后续将定期跟踪整改进度，确保各项整改措施按计划推进落实。",
        "运营专班",
    ]
    level1 = [paragraph for paragraph in non_empty if paragraph.text.strip().startswith(("一、", "二、", "三、"))]
    level2 = [paragraph for paragraph in non_empty if paragraph.text.strip().startswith(("（一）", "（二）", "（三）"))]
    assert {paragraph.runs[0].font.name for paragraph in level1} == {"黑体"}
    assert {paragraph.paragraph_format.first_line_indent.pt for paragraph in level1} == {32}
    assert {paragraph.runs[0].font.name for paragraph in level2} == {"楷体_GB2312"}
    assert non_empty[-1].runs[0].font.name == "仿宋_GB2312"


def test_management_method_single_paragraph_uses_scripted_chapter_recovery(tmp_path: Path) -> None:
    input_path = tmp_path / "management_method_source.docx"
    source = Document()
    source.add_paragraph(
        "超储物资内部调剂消耗指引"
        "超储物资是指库龄通常在一年以内，因项目结余、采购计划变更或储备定额调整而形成的，库存数量超出实际需求但物理性能完好、具备完整使用价值的备品配件及材料。"
        "基层企业应通过 ERP 系统或物资管理模块，定期筛选“无动态、高库存”物资，将其列入超储物资调剂清单进行专项管理。"
        "集团供应链管理部作为归口管理部门，负责制定管理制度，建立信息平台，并对各单位调剂完成情况进行绩效考核。"
        "集团物资供应中心（调剂中心）负责跨二级单位调剂的交易撮合、物流协调及华能商城调剂模块的日常运营。"
        "二级单位（管理中心）负责本单位内部超储物资的认定、价格审批及所属基层企业的利库监督。"
        "基层企业（执行主体）负责实物信息的实时发布、质量维护、实物交接及账务处理。"
        "超储物资上架信息须包含物资编码、参数规格、交易价格、实物照片及技术说明书等关键要素。"
        "需求单位在发起采购申请前，系统将自动匹配集团范围内超储物资库。"
        "超储物资的调剂价格主要依据交易双方的股权性质及资产账面价值确定。"
        "资产减值是指资产在调剂时的可收回金额低于其账面余额的差额。"
        "资产评估是通过专业机构评定确定调剂资产在特定时点的公允价值。"
        "重置完全价是指在当前市场环境下，重新购置与该物资完全相同的新品所需的全部成本。"
        "超储物资调剂产生的装卸、运输及保险费用原则上由使用方（需求单位）承担。"
        "为提高资源流转效率，集团公司授权物资供应中心对账面价值为零但仍有使用价值的超储物资，行使跨二级单位调拨的快速审批权。"
        "所有调剂业务必须在华能商城“联储联备专区”完成从发布、下单、发货到确认验收的全流程线上闭环操作。"
        "财务处理与税务合规方面，调出方产生的调剂收入计入“其他业务收入”或“营业外收入”。"
        "超储物资调剂适用“现状交付、风险自担”原则。"
        "考核激励层面，若超储物资调剂导致调出方产生资产损失并降低利润，在年度考核计算时，该部分损失不计入当年利润指标考核计算值。"
    )
    source.save(str(input_path))
    output_path = tmp_path / "management_method_formatted.docx"

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--assume-detected-type")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    non_empty = [paragraph for paragraph in output.paragraphs if paragraph.text.strip()]
    texts = [paragraph.text.strip() for paragraph in non_empty]
    expected_chapters = [
        "一、管理职责分工",
        "二、上架信息发布要求",
        "三、需求匹配与优先调剂",
        "四、调剂定价规则",
        "五、资产减值处理",
        "六、资产评估机制",
        "七、重置完全价核定",
        "八、费用承担",
        "九、零值物资快速调拨",
        "十、线上操作流程",
        "十一、财务处理与税务合规",
        "十二、交付验收与质保",
        "十三、考核激励",
    ]
    assert texts[0] == "超储物资内部调剂消耗指引"
    assert all(chapter in texts for chapter in expected_chapters)
    assert not any(text in {"职责分工", "信息发布", "调剂价格", "考核激励"} for text in texts)
    level1 = [paragraph for paragraph in non_empty if paragraph.text.strip() in expected_chapters]
    assert len(level1) == len(expected_chapters)
    assert {paragraph.runs[0].font.name for paragraph in level1} == {"黑体"}
    assert {paragraph.paragraph_format.first_line_indent.pt for paragraph in level1} == {32}


def test_management_method_real_single_paragraph_fixture_recovers_structure(tmp_path: Path) -> None:
    fixture_path = Path("/Users/liuzigeng/Ageng的自媒体/公文写作项目/xiaohongshu/对比图/管理办法.docx")
    assert fixture_path.exists()

    entry_modes = [
        ("assume_detected", ["--assume-detected-type"]),
        ("generic_flag", ["--generic-formal-text"]),
        ("generic_doc_type", ["--doc-type", "通用正式文本"]),
    ]
    expected_fragments = [
        "超储物资内部调剂消耗指引",
        "调剂定价规则",
        "财务处理与税务合规",
        "考核激励",
    ]

    for name, extra_args in entry_modes:
        input_path = tmp_path / f"{name}_source.docx"
        shutil.copyfile(fixture_path, input_path)
        output_path = tmp_path / f"{name}_formatted.docx"
        report_path = output_path.with_suffix(".report.json")

        result = run_format_cli(str(input_path), "-o", str(output_path), "--report", *extra_args)

        assert result.returncode == 0, result.stderr
        output = Document(str(output_path))
        texts = [paragraph.text.strip() for paragraph in output.paragraphs if paragraph.text.strip()]
        assert len(texts) > 10
        assert texts[0] == "超储物资内部调剂消耗指引"
        assert all(any(fragment in text for text in texts) for fragment in expected_fragments)
        assert not (len(texts) == 1 and len(texts[0]) > 1000)

        report = json.loads(report_path.read_text(encoding="utf-8"))
        recovery_ops = [operation for operation in report["operations"] if operation["kind"] == "chapter_recovery"]
        assert recovery_ops
        assert recovery_ops[0]["params"]["method"] == "management_method"


def test_generic_formal_text_flag_overrides_standard_spec_auto_detection(tmp_path: Path) -> None:
    input_path = tmp_path / "standard_like_without_toc.docx"
    source = Document()
    for text in [
        "发电企业供应链服务标准",
        "前    言",
        "本文件说明相关服务要求。",
        "1  范围",
        "本文件规定了供应链服务范围。",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "standard_like_as_generic.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generic-formal-text")

    assert result.returncode == 0, result.stderr
    assert output_path.exists()
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["doc_type"] == "通用正式文本"


def test_generate_toc_flag_adds_word_toc_field_for_clear_headings(tmp_path: Path) -> None:
    input_path = tmp_path / "report_with_headings.docx"
    source = Document()
    for text in [
        "关于安全生产整改情况的报告",
        "集团公司：",
        "一、总体情况",
        "本段介绍总体情况。",
        "（一）隐患排查情况",
        "本段介绍排查情况。",
        "1. 重点问题",
        "本段介绍重点问题。",
        "二、下一步工作",
        "本段介绍下一步工作。",
        "某某公司",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "report_with_toc.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generate-toc")

    assert result.returncode == 0, result.stderr
    xml = document_xml(output_path)
    assert "TOC" in xml
    assert "fldChar" in xml
    assert "outlineLvl" in xml

    report = json.loads(report_path.read_text(encoding="utf-8"))
    toc_ops = [operation for operation in report["operations"] if operation["kind"] == "toc_generation"]
    assert toc_ops
    assert toc_ops[0]["params"]["action"] == "generated"
    assert toc_ops[0]["params"]["heading_count"] >= 4


def test_generate_toc_flag_skips_when_headings_are_not_clear(tmp_path: Path) -> None:
    input_path = tmp_path / "short_notice.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "short_notice_no_toc.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--generate-toc")

    assert result.returncode == 0, result.stderr
    assert "TOC" not in document_xml(output_path)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    toc_ops = [operation for operation in report["operations"] if operation["kind"] == "toc_generation"]
    assert toc_ops
    assert toc_ops[0]["params"]["action"] == "skipped_unclear_headings"
    assert toc_ops[0]["params"]["heading_count"] == 0


def test_format_imprint_flag_formats_existing_imprint_and_reports_action(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_with_imprint.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
        "某某办公室                        2026年6月21日印发",
    ]:
        paragraph = source.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = None
    source.save(str(input_path))

    output_path = tmp_path / "notice_with_imprint_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-imprint")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    imprint = next(paragraph for paragraph in output.paragraphs if "印发" in paragraph.text)
    assert imprint.paragraph_format.first_line_indent.pt == 0
    assert imprint.runs[0].font.name == "仿宋_GB2312"
    assert imprint.runs[0].font.size.pt == 16

    report = json.loads(report_path.read_text(encoding="utf-8"))
    imprint_ops = [operation for operation in report["operations"] if operation["kind"] == "imprint_format"]
    assert imprint_ops
    assert imprint_ops[0]["params"]["action"] == "formatted"
    assert imprint_ops[0]["params"]["imprint_count"] == 1


def test_format_imprint_flag_reports_no_imprint_without_adding_one(tmp_path: Path) -> None:
    input_path = tmp_path / "notice_without_imprint.docx"
    source = Document()
    for text in [
        "关于开展安全检查的通知",
        "各部门：",
        "请按要求开展安全检查。",
        "某某办公室",
        "2026年6月21日",
    ]:
        source.add_paragraph(text)
    source.save(str(input_path))

    output_path = tmp_path / "notice_without_imprint_formatted.docx"
    report_path = output_path.with_suffix(".report.json")

    result = run_format_cli(str(input_path), "-o", str(output_path), "--report", "--format-imprint")

    assert result.returncode == 0, result.stderr
    output = Document(str(output_path))
    assert not any("印发" in paragraph.text for paragraph in output.paragraphs)

    report = json.loads(report_path.read_text(encoding="utf-8"))
    imprint_ops = [operation for operation in report["operations"] if operation["kind"] == "imprint_format"]
    assert imprint_ops
    assert imprint_ops[0]["params"]["action"] == "no_imprint"
    assert imprint_ops[0]["params"]["imprint_count"] == 0
