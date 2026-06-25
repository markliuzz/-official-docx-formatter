#!/usr/bin/env python3
"""Format Chinese official DOCX documents using JSON profiles."""

from __future__ import annotations

import argparse
import copy
import json
import re
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile
from typing import Any, Dict, Iterable, Optional, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - environment message
    raise SystemExit("python-docx is required. Install with: pip install python-docx") from exc

from official_docx_engine.models import FormatOperation
from official_docx_engine.models import DocumentSnapshot, ParagraphSnapshot
from official_docx_engine.imprint import format_existing_imprint
from official_docx_engine.page_numbers import apply_page_numbers, inspect_footers
from official_docx_engine.standard_text import build_standard_text_document, looks_like_standard_text
from official_docx_engine.tables import append_and_format_source_tables, copy_and_format_table
from official_docx_engine.toc import generate_toc_if_clear

SKILL_DIR = Path(__file__).resolve().parents[1]
PROFILES_DIR = SKILL_DIR / "profiles"
DOCUMENT_TYPES_FILE = SKILL_DIR / "references" / "document_types.json"
BASE_PROFILE = "standard-party-government"
GENERIC_FORMAL_TEXT = "通用正式文本"
STANDARD_SPEC_TEXT = "标准规范文本"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_VAL = f"{{{W_NS}}}val"
W_ABSTRACT_NUM_ID = f"{{{W_NS}}}abstractNumId"
W_ILVL = f"{{{W_NS}}}ilvl"
W_NUM_ID = f"{{{W_NS}}}numId"

CN_DATE_RE = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")
SENTENCE_PUNCT_RE = re.compile(r"[。；;！？!?]")
NUMBERED_ITEM_RE = re.compile(
    r"^([一二三四五六七八九十]+[、，,]|\d+[.．、]|\d+\s|（[一二三四五六七八九十\d]+）|\([一二三四五六七八九十\d]+\)|[a-z]）)"
)
GLUED_TITLE_KEYWORD_RE = re.compile(r"(通知|通报|报告|请示|批复|函|纪要|意见|决定|决议|公告|通告|方案|总结|汇报)")
GLUED_SECTION_HEADINGS = (
    "存在的问题",
    "主要问题",
    "原因分析",
    "解决措施",
    "整改措施",
    "整改情况",
    "时间计划",
    "工作计划",
    "下一步工作",
    "下一步安排",
    "基本情况",
    "总体情况",
    "有关情况",
    "处理建议",
    "工作要求",
)
UNNUMBERED_HEADING_SUFFIXES = (
    "情况",
    "问题",
    "措施",
    "计划",
    "安排",
    "要求",
    "建议",
    "成效",
    "背景",
    "目标",
    "原则",
    "范围",
    "原因",
    "风险",
    "结论",
    "说明",
)
SIGNATURE_HINT_RE = re.compile(r"(公司|集团|局|厅|部|委|办|处|科|院|中心|办公室|专班|小组|委员会)$")
SPACED_SUBHEADING_SUFFIXES = ("规范", "异常", "图片", "方案", "问题", "措施")
MANAGEMENT_METHOD_TITLE = "超储物资内部调剂消耗指引"
MANAGEMENT_METHOD_CHAPTERS = (
    ("管理职责分工", ("集团供应链管理部",)),
    ("上架信息发布要求", ("超储物资上架信息", "上架信息须")),
    ("需求匹配与优先调剂", ("需求单位在发起",)),
    ("调剂定价规则", ("超储物资的调剂价格",)),
    ("资产减值处理", ("资产减值是指", "资产减值，定义")),
    ("资产评估机制", ("资产评估是", "资产评估，定义")),
    ("重置完全价核定", ("重置完全价是指", "重置完全价，定义")),
    ("费用承担", ("超储物资调剂产生", "费用主体：超储物资调剂")),
    ("零值物资快速调拨", ("为提高资源流转效率",)),
    ("线上操作流程", ("所有调剂业务必须", "线上化闭环")),
    ("财务处理与税务合规", ("财务处理与税务合规方面", "调出方处理：")),
    ("交付验收与质保", ("超储物资调剂适用", "现状交付机制")),
    ("考核激励", ("考核激励层面", "考核激励机制")),
)


def _engine_imports():
    from official_docx_engine.diagnostics import diagnose_snapshot
    from official_docx_engine.docx_reader import read_docx_snapshot
    from official_docx_engine.format_plan import build_format_plan
    from official_docx_engine.reporting import write_report_json
    from official_docx_engine.structure import analyze_structure

    return read_docx_snapshot, analyze_structure, diagnose_snapshot, build_format_plan, write_report_json


def deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    result = copy.deepcopy(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = copy.deepcopy(value)
    return result


def load_profile(profile_id: str) -> Dict[str, Any]:
    path = PROFILES_DIR / f"{profile_id}.json"
    if not path.exists():
        raise SystemExit(f"profile not found: {profile_id}")
    profile = json.loads(path.read_text(encoding="utf-8"))
    parent_id = profile.get("inherits")
    if parent_id:
        parent = load_profile(parent_id)
        profile = deep_merge(parent, profile)
    return profile


def load_document_types() -> Dict[str, Any]:
    return json.loads(DOCUMENT_TYPES_FILE.read_text(encoding="utf-8"))


def classify_lines_for_type(lines: list[str]) -> Dict[str, Any]:
    import sys

    scripts_dir = str(Path(__file__).resolve().parent)
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from classify_document import classify_lines

    return classify_lines(lines)


def normalize_line(text: str, enabled: bool, space_mode: str) -> str:
    if not enabled:
        return text
    import sys

    scripts_dir = str(Path(__file__).resolve().parent)
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    from normalize_text import normalize_text

    return normalize_text(text, space_mode=space_mode)


def split_glued_single_paragraph(lines: list[str]) -> list[str]:
    """Conservatively split an extreme one-paragraph official document draft."""

    non_empty = [line.strip() for line in lines if line.strip()]
    if len(non_empty) != 1:
        return lines
    text = non_empty[0]
    management_blocks = _split_management_method(text)
    if management_blocks is not None:
        return management_blocks
    if len(text) < 120 or SENTENCE_PUNCT_RE.search(text[:40]):
        return lines

    title_match = GLUED_TITLE_KEYWORD_RE.search(text[:80])
    if not title_match:
        return lines
    title = text[: title_match.end()].strip()
    rest = text[title_match.end() :].strip()
    if not title or not rest:
        return lines

    blocks = [title]
    blocks.extend(_split_glued_body(rest))
    return blocks if len(blocks) > 2 else lines


def _split_management_method(text: str) -> list[str] | None:
    compact = text.strip()
    if not compact.startswith(MANAGEMENT_METHOD_TITLE):
        return None

    chapter_points: list[tuple[int, str, str]] = []
    for title, anchors in MANAGEMENT_METHOD_CHAPTERS:
        match = _find_first_anchor(compact, anchors)
        if match is None:
            return None
        position, anchor = match
        chapter_points.append((position, title, anchor))
    chapter_points.sort(key=lambda item: item[0])

    if [title for _, title, _ in chapter_points] != [title for title, _ in MANAGEMENT_METHOD_CHAPTERS]:
        return None

    blocks = [MANAGEMENT_METHOD_TITLE]
    lead = compact[len(MANAGEMENT_METHOD_TITLE) : chapter_points[0][0]].strip()
    if lead:
        blocks.extend(_split_sentence_blocks(lead))

    for index, (position, title, _anchor) in enumerate(chapter_points, 1):
        body_end = chapter_points[index][0] if index < len(chapter_points) else len(compact)
        body = compact[position:body_end].strip()
        blocks.append(f"{_cn_number(index)}、{title}")
        blocks.extend(_split_sentence_blocks(body))
    return blocks if len(blocks) > 4 else None


def _find_first_anchor(text: str, anchors: tuple[str, ...]) -> Optional[tuple[int, str]]:
    matches = [(position, anchor) for anchor in anchors if (position := text.find(anchor)) != -1]
    if not matches:
        return None
    return min(matches, key=lambda item: item[0])


def _split_sentence_blocks(text: str) -> list[str]:
    stripped = text.strip()
    if not stripped:
        return []
    pieces = [piece.strip() for piece in re.split(r"(?<=[。；;])", stripped) if piece.strip()]
    return pieces or [stripped]


def _split_glued_body(text: str) -> list[str]:
    spaced_blocks = _split_spaced_glued_body(text)
    if spaced_blocks is not None:
        return spaced_blocks

    heading_positions: list[tuple[int, str]] = []
    for heading in GLUED_SECTION_HEADINGS:
        start = 0
        while True:
            position = text.find(heading, start)
            if position == -1:
                break
            previous = text[position - 1] if position > 0 else ""
            if position == 0 or previous in "。；;！？!?":
                heading_positions.append((position, heading))
            start = position + len(heading)

    heading_positions = sorted(set(heading_positions))
    if not heading_positions:
        return _split_trailing_signature(text)

    blocks: list[str] = []
    cursor = 0
    for index, (position, heading) in enumerate(heading_positions):
        prefix = text[cursor:position].strip()
        if prefix:
            blocks.extend(_split_trailing_signature(prefix))
        blocks.append(heading)
        content_start = position + len(heading)
        content_end = heading_positions[index + 1][0] if index + 1 < len(heading_positions) else len(text)
        content = text[content_start:content_end].strip()
        if content:
            blocks.extend(_split_trailing_signature(content))
        cursor = content_end

    suffix = text[cursor:].strip()
    if suffix:
        blocks.extend(_split_trailing_signature(suffix))
    return blocks


def _split_trailing_signature(text: str) -> list[str]:
    stripped = text.strip()
    if not stripped:
        return []
    for index in range(len(stripped) - 1, -1, -1):
        if stripped[index] not in "。；;！？!?":
            continue
        tail = stripped[index + 1 :].strip()
        head = stripped[: index + 1].strip()
        if head and 2 <= len(tail) <= 40 and not SENTENCE_PUNCT_RE.search(tail) and SIGNATURE_HINT_RE.search(tail):
            return [head, tail]
        break
    return [stripped]


def _split_spaced_glued_body(text: str) -> list[str] | None:
    segments = [segment.strip() for segment in re.split(r"\s+", text) if segment.strip()]
    main_heading_count = sum(1 for segment in segments if segment in GLUED_SECTION_HEADINGS)
    short_structure_count = sum(1 for segment in segments if _is_spaced_subheading_candidate(segment))
    if len(segments) < 6 or main_heading_count < 2 or short_structure_count < 2:
        return None

    blocks: list[str] = []
    current_main = ""
    main_count = 0
    sub_count = 0
    for segment in segments:
        if segment in GLUED_SECTION_HEADINGS:
            main_count += 1
            sub_count = 0
            current_main = segment
            blocks.append(f"{_cn_number(main_count)}、{segment}")
            continue
        if current_main and current_main != "时间计划" and _is_spaced_subheading_candidate(segment):
            sub_count += 1
            blocks.append(f"（{_cn_number(sub_count)}）{segment}")
            continue
        blocks.extend(_split_trailing_signature(segment))
    return blocks


def _is_spaced_subheading_candidate(text: str) -> bool:
    compact = re.sub(r"\s+", "", text.strip())
    if (
        not compact
        or len(compact) > 24
        or SENTENCE_PUNCT_RE.search(compact)
        or NUMBERED_ITEM_RE.match(compact)
        or compact in GLUED_SECTION_HEADINGS
        or SIGNATURE_HINT_RE.search(compact)
    ):
        return False
    return compact.endswith(SPACED_SUBHEADING_SUFFIXES)


def repair_glued_snapshot(snapshot: DocumentSnapshot) -> DocumentSnapshot:
    lines = [paragraph.text for paragraph in snapshot.non_empty_paragraphs]
    repaired = split_glued_single_paragraph(lines)
    if repaired == lines:
        return snapshot
    paragraphs = tuple(
        ParagraphSnapshot(
            index=index,
            text=text,
            style_name=snapshot.non_empty_paragraphs[0].style_name,
            alignment=snapshot.non_empty_paragraphs[0].alignment,
            runs=snapshot.non_empty_paragraphs[0].runs if index == 0 else (),
        )
        for index, text in enumerate(repaired)
    )
    return DocumentSnapshot(path=snapshot.path, paragraphs=paragraphs, table_count=snapshot.table_count)


def build_skeleton_body(doc_type: str) -> list[str]:
    document_types = load_document_types()
    if doc_type not in document_types:
        supported = "、".join(document_types)
        raise SystemExit(f"unsupported doc type: {doc_type}; supported: {supported}")

    info = document_types[doc_type]
    body: list[str] = []
    for index, section in enumerate(info.get("sections", []), 1):
        body.append(f"{_cn_number(index)}、{section}")
        body.append(f"请补充{section}。")
    ending = info.get("ending")
    if ending:
        body.append(ending)
    return body


def _cn_number(value: int) -> str:
    numbers = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 0 <= value <= 10:
        return numbers[value]
    if value < 20:
        return f"十{numbers[value - 10]}"
    if value < 100:
        tens, ones = divmod(value, 10)
        return f"{numbers[tens]}十{numbers[ones] if ones else ''}"
    return str(value)


def _cn_counting_number(value: int) -> str:
    digits = "零一二三四五六七八九"
    if value <= 0:
        return str(value)
    if value <= 10:
        return "十" if value == 10 else digits[value]
    if value < 20:
        return "十" + digits[value % 10]
    if value < 100:
        tens, ones = divmod(value, 10)
        return digits[tens] + "十" + (digits[ones] if ones else "")
    return str(value)


def load_numbering_formats(input_path: Path) -> dict[tuple[str, str], dict[str, str]]:
    """Read Word numbering definitions needed to materialize visible list prefixes."""

    try:
        with ZipFile(input_path) as archive:
            if "word/numbering.xml" not in archive.namelist():
                return {}
            root = ET.fromstring(archive.read("word/numbering.xml"))
    except Exception:
        return {}

    ns = {"w": W_NS}
    abstract_levels: dict[tuple[str, str], dict[str, str]] = {}
    for abstract_num in root.findall("w:abstractNum", ns):
        abstract_id = abstract_num.get(W_ABSTRACT_NUM_ID)
        if abstract_id is None:
            continue
        for level in abstract_num.findall("w:lvl", ns):
            ilvl = level.get(W_ILVL, "0")
            num_fmt = level.find("w:numFmt", ns)
            lvl_text = level.find("w:lvlText", ns)
            start = level.find("w:start", ns)
            abstract_levels[(abstract_id, ilvl)] = {
                "num_fmt": num_fmt.get(W_VAL, "decimal") if num_fmt is not None else "decimal",
                "lvl_text": lvl_text.get(W_VAL, "%1") if lvl_text is not None else "%1",
                "start": start.get(W_VAL, "1") if start is not None else "1",
            }

    formats: dict[tuple[str, str], dict[str, str]] = {}
    for num in root.findall("w:num", ns):
        num_id = num.get(W_NUM_ID)
        abstract_num_id = num.find("w:abstractNumId", ns)
        if num_id is None or abstract_num_id is None:
            continue
        abstract_id = abstract_num_id.get(W_VAL)
        for (candidate_abstract_id, ilvl), info in abstract_levels.items():
            if candidate_abstract_id == abstract_id:
                formats[(num_id, ilvl)] = info
    return formats


def visible_paragraph_text(paragraph: Paragraph, numbering_formats: dict[tuple[str, str], dict[str, str]], counters: dict[tuple[str, str], int]) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    numbering_prefix = _paragraph_numbering_prefix(paragraph, numbering_formats, counters)
    if numbering_prefix and not text.startswith(numbering_prefix):
        return numbering_prefix + text
    return text


def _paragraph_numbering_prefix(paragraph: Paragraph, numbering_formats: dict[tuple[str, str], dict[str, str]], counters: dict[tuple[str, str], int]) -> str:
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is None:
        return ""
    numbering = paragraph_properties.find(qn("w:numPr"))
    if numbering is None:
        return ""
    num_id_el = numbering.find(qn("w:numId"))
    if num_id_el is None:
        return ""
    ilvl_el = numbering.find(qn("w:ilvl"))
    num_id = num_id_el.get(qn("w:val"))
    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
    if num_id is None:
        return ""
    info = numbering_formats.get((num_id, ilvl))
    if info is None:
        return ""

    key = (num_id, ilvl)
    start = int(info.get("start", "1"))
    current = counters.get(key, start - 1) + 1
    counters[key] = current
    number_text = _format_number(current, info.get("num_fmt", "decimal"))
    return info.get("lvl_text", "%1").replace("%1", number_text)


def _format_number(value: int, num_fmt: str) -> str:
    if num_fmt in {"chineseCounting", "chineseCountingThousand"}:
        return _cn_counting_number(value)
    return str(value)


def preferred_font(profile: Dict[str, Any], key: str) -> str:
    fonts = profile.get("fonts", {}).get(key, {})
    fallbacks = fonts.get("fallbacks") or []
    return fallbacks[0] if fallbacks else "仿宋"


def font_size(profile: Dict[str, Any], key: str, default: float = 16) -> float:
    return float(profile.get("fonts", {}).get(key, {}).get("size_pt", default))


def font_bold(profile: Dict[str, Any], key: str) -> bool:
    return bool(profile.get("fonts", {}).get(key, {}).get("bold", False))


def add_run(paragraph, text: str, font: str, size_pt: float, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_common_paragraph_format(paragraph, profile: Dict[str, Any], first_line: bool = True) -> None:
    layout = profile.get("layout", {})
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(float(layout.get("line_spacing_pt", 28)))
    paragraph.paragraph_format.space_before = Pt(float(layout.get("space_before_pt", 0)))
    paragraph.paragraph_format.space_after = Pt(float(layout.get("space_after_pt", 0)))
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(32)
    else:
        paragraph.paragraph_format.first_line_indent = Pt(0)


def setup_page(doc: Document, profile: Dict[str, Any]) -> None:
    page = profile.get("page", {})
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(float(page.get("top_margin_cm", 3.7)))
    section.bottom_margin = Cm(float(page.get("bottom_margin_cm", 3.5)))
    section.left_margin = Cm(float(page.get("left_margin_cm", 2.8)))
    section.right_margin = Cm(float(page.get("right_margin_cm", 2.6)))


def looks_like_inline_body(text: str) -> bool:
    stripped = text.strip()
    if re.search(r"[。；;：:]", stripped):
        return True
    return len(stripped) > 34


def unnumbered_heading_key(text: str) -> Optional[str]:
    stripped = text.strip()
    compact = re.sub(r"\s+", "", stripped)
    if (
        not compact
        or len(compact) > 24
        or SENTENCE_PUNCT_RE.search(stripped)
        or NUMBERED_ITEM_RE.match(stripped)
        or CN_DATE_RE.match(compact)
        or SIGNATURE_HINT_RE.search(compact)
    ):
        return None
    if compact in GLUED_SECTION_HEADINGS or compact.endswith(UNNUMBERED_HEADING_SUFFIXES):
        return "level1"
    return None


def hierarchy_key(text: str) -> Optional[str]:
    if re.match(r"^[一二三四五六七八九十]+[、，,]", text):
        return "level1"
    if re.match(r"^(（[一二三四五六七八九十]）|\([一二三四五六七八九十]\))", text):
        return None if looks_like_inline_body(text) else "level2"
    if re.match(r"^\d+[\.．]", text):
        return None if looks_like_inline_body(text) else "body"
    if re.match(r"^(（\d+）|\(\d+\))", text):
        return None if looks_like_inline_body(text) else "body"
    return unnumbered_heading_key(text)


def split_source_document(input_path: Path, normalize: bool = True, space_mode: str = "keep_en_boundary") -> Tuple[str, str, list[str], Optional[str], Optional[str]]:
    source = Document(str(input_path))
    paragraphs = [
        normalize_line(p.text.strip(), normalize, space_mode)
        for p in source.paragraphs
        if p.text.strip()
    ]
    title = paragraphs[0] if paragraphs else ""
    recipient = ""
    body_start = 1

    for index, text in enumerate(paragraphs[1:], 1):
        if text.endswith(("：", ":")) and len(text) <= 40:
            recipient = text.rstrip("：:").strip()
            body_start = index + 1
            break

    body = paragraphs[body_start:]
    issuer = None
    date_text = None
    if len(body) >= 2 and CN_DATE_RE.match(body[-1]):
        issuer = body[-2]
        date_text = body[-1]
        body = body[:-2]
    return title, recipient, body, issuer, date_text


def recipient_candidate(text: str) -> bool:
    stripped = text.strip()
    compact = re.sub(r"\s+", "", stripped)
    return (
        stripped.endswith(("：", ":"))
        and len(compact) <= 50
        and not SENTENCE_PUNCT_RE.search(stripped)
        and not NUMBERED_ITEM_RE.match(stripped)
    )


TITLE_CONTINUATION_RE = re.compile(r"(通知|通报|报告|请示|批复|函|纪要|意见|决定|决议|公告|通告|方案|总结|汇报|情况|说明|手册|标准|规范|办法|指引|流程|审批单|规格说明书)")
PARENTHETICAL_TITLE_RE = re.compile(r"^（[^）]{1,20}）$")


def title_continuation_candidate(text: str, title_lines: list[str]) -> bool:
    stripped = text.strip()
    compact = re.sub(r"\s+", "", stripped)
    if (
        not compact
        or len(compact) > 40
        or recipient_candidate(stripped)
        or SENTENCE_PUNCT_RE.search(stripped)
        or NUMBERED_ITEM_RE.match(stripped)
    ):
        return False
    return bool(TITLE_CONTINUATION_RE.search(compact) or (title_lines and PARENTHETICAL_TITLE_RE.match(compact)))


def split_source_paragraphs(paragraphs: list[str]) -> Tuple[list[str], str, int, Optional[str], Optional[str]]:
    title_lines = paragraphs[:1]
    search_end = min(len(paragraphs), 5)

    for index in range(1, search_end):
        if title_continuation_candidate(paragraphs[index], title_lines):
            title_lines.append(paragraphs[index])
            continue
        break

    recipient = ""
    body_start = len(title_lines)
    for index in range(body_start, search_end):
        text = paragraphs[index]
        if recipient_candidate(text):
            recipient = text.rstrip("：:").strip()
            body_start = index + 1
            break

    issuer = None
    date_text = None
    if len(paragraphs) - body_start >= 2 and CN_DATE_RE.match(paragraphs[-1]):
        issuer = paragraphs[-2]
        date_text = paragraphs[-1]
    return title_lines, recipient, body_start, issuer, date_text


def add_title(doc: Document, title: str, profile: Dict[str, Any], space_after_pt: float = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    add_run(paragraph, title, preferred_font(profile, "title"), font_size(profile, "title", 22), font_bold(profile, "title"))


def add_recipient(doc: Document, recipient: str, profile: Dict[str, Any]) -> None:
    if not recipient:
        return
    paragraph = doc.add_paragraph()
    set_common_paragraph_format(paragraph, profile, first_line=False)
    add_run(paragraph, f"{recipient}：", preferred_font(profile, "body"), font_size(profile, "body", 16), False)


def add_body_paragraph(doc: Document, text: str, profile: Dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    key = hierarchy_key(text) or "body"
    set_common_paragraph_format(paragraph, profile, first_line=True)
    add_run(paragraph, text, preferred_font(profile, key), font_size(profile, key, 16), font_bold(profile, key))


def add_footer(doc: Document, issuer: Optional[str], date_text: Optional[str], profile: Dict[str, Any]) -> None:
    if not issuer and not date_text:
        return
    signature = profile.get("signature", {})
    blank_lines_before = int(signature.get("blank_lines_before", 1))
    right_empty_chars = float(signature.get("right_empty_chars", 2))
    body_font = preferred_font(profile, "body")
    body_size = font_size(profile, "body", 16)
    for _ in range(max(blank_lines_before, 0)):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = Pt(float(profile.get("layout", {}).get("line_spacing_pt", 28)))
        add_run(paragraph, "", body_font, body_size, False)
    for value in [issuer, date_text]:
        if not value:
            continue
        paragraph = doc.add_paragraph()
        set_common_paragraph_format(paragraph, profile, first_line=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.right_indent = Pt(body_size * right_empty_chars)
        add_run(paragraph, value, body_font, body_size, False)


def build_document(title: str, recipient: str, body: Iterable[str], issuer: Optional[str], date_text: Optional[str], profile: Dict[str, Any], normalize: bool = False, space_mode: str = "keep_en_boundary") -> Document:
    doc = Document()
    setup_page(doc, profile)
    title = normalize_line(title, normalize, space_mode)
    recipient = normalize_line(recipient, normalize, space_mode)
    issuer = normalize_line(issuer, normalize, space_mode) if issuer else issuer
    date_text = normalize_line(date_text, normalize, space_mode) if date_text else date_text
    if title:
        add_title(doc, title, profile)
    add_recipient(doc, recipient, profile)
    for text in body:
        text = normalize_line(text.strip(), normalize, space_mode)
        if text:
            add_body_paragraph(doc, text, profile)
    add_footer(doc, issuer, date_text, profile)
    return doc


def build_document_from_source(
    input_path: Path,
    title_override: Optional[str],
    recipient_override: str,
    issuer_override: Optional[str],
    date_override: Optional[str],
    profile: Dict[str, Any],
    normalize: bool,
    space_mode: str,
    generic_formal_text: bool = False,
) -> tuple[Document, int, Optional[str]]:
    source = Document(str(input_path))
    numbering_formats = load_numbering_formats(input_path)
    paragraph_text_counters: dict[tuple[str, str], int] = {}
    paragraph_texts = [
        visible_paragraph_text(paragraph, numbering_formats, paragraph_text_counters)
        for paragraph in source.paragraphs
        if paragraph.text.strip()
    ]
    repaired_paragraph_texts = split_glued_single_paragraph(paragraph_texts)
    glued_repaired = repaired_paragraph_texts != paragraph_texts
    recovery_method = None
    if glued_repaired:
        recovery_method = "management_method" if paragraph_texts[0].startswith(MANAGEMENT_METHOD_TITLE) else "glued_single_paragraph"
    paragraph_texts = repaired_paragraph_texts
    title_lines, detected_recipient, body_start, detected_issuer, detected_date = split_source_paragraphs(paragraph_texts)
    if generic_formal_text:
        body_start = len(title_lines)
        recipient = recipient_override
        issuer = issuer_override
        date_text = date_override
        detected_issuer = None
        detected_date = None
    else:
        recipient = recipient_override or detected_recipient
        issuer = issuer_override or detected_issuer
        date_text = date_override or detected_date

    doc = Document()
    setup_page(doc, profile)
    if title_override:
        add_title(doc, normalize_line(title_override, normalize, space_mode), profile)
    else:
        for index, title in enumerate(title_lines):
            add_title(
                doc,
                normalize_line(title, normalize, space_mode),
                profile,
                space_after_pt=12 if index == len(title_lines) - 1 else 0,
            )
    add_recipient(doc, normalize_line(recipient, normalize, space_mode), profile)

    if glued_repaired:
        for text in paragraph_texts[body_start:]:
            if detected_issuer and detected_date and text in {detected_issuer, detected_date}:
                continue
            add_body_paragraph(doc, normalize_line(text, normalize, space_mode), profile)
        add_footer(doc, issuer, date_text, profile)
        return doc, 0, recovery_method

    paragraph_ordinal = -1
    copied_tables = 0
    source_table_index = 0
    body_text_counters: dict[tuple[str, str], int] = {}
    for child in source.element.body.iterchildren():
        if child.tag == qn("w:p"):
            source_paragraph = Paragraph(child, source)
            text = visible_paragraph_text(source_paragraph, numbering_formats, body_text_counters)
            if not text:
                continue
            paragraph_ordinal += 1
            if paragraph_ordinal < body_start:
                continue
            if paragraph_ordinal >= len(paragraph_texts) - 2 and detected_issuer and detected_date and text in {detected_issuer, detected_date}:
                continue
            normalized = normalize_line(text, normalize, space_mode)
            if normalized:
                add_body_paragraph(doc, normalized, profile)
        elif child.tag == qn("w:tbl") and paragraph_ordinal >= body_start - 1:
            if source_table_index < len(source.tables):
                copy_and_format_table(doc, source.tables[source_table_index], profile)
                copied_tables += 1
            source_table_index += 1
        elif child.tag == qn("w:tbl"):
            source_table_index += 1

    add_footer(doc, issuer, date_text, profile)
    return doc, copied_tables, recovery_method


def smoke_check(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return f"paragraphs={len(paragraphs)}"


def _report_path(output_path: Path, report_arg: Optional[str]) -> Optional[Path]:
    if report_arg in {None, ""}:
        return output_path.with_suffix(".report.json")
    return Path(report_arg)


def main() -> int:
    parser = argparse.ArgumentParser(description="Format Chinese official DOCX documents")
    parser.add_argument("input", nargs="?", help="Input .docx path. Omit when using --text-file.")
    parser.add_argument("-o", "--output", required=True, help="Output .docx path")
    parser.add_argument("--profile", default=BASE_PROFILE, help=argparse.SUPPRESS)
    parser.add_argument("--title", help="Title when creating from text")
    parser.add_argument("--recipient", default="", help="Recipient when creating from text")
    parser.add_argument("--issuer", help="Issuer override")
    parser.add_argument("--date", dest="date_text", help="Date override, e.g. 2026年6月19日")
    parser.add_argument("--text-file", help="Create a formatted .docx from a plain text file")
    parser.add_argument("--doc-type", help="Document type, e.g. 请示, 报告, 通知, 函, 纪要, 通用正式文本")
    parser.add_argument("--create-skeleton", action="store_true", help="Create a document skeleton from --doc-type")
    parser.add_argument("--assume-detected-type", action="store_true", help="Continue when classifier is confident; otherwise stop and ask user")
    parser.add_argument(
        "--report",
        nargs="?",
        const="",
        default=None,
        help="Write a JSON formatting report for DOCX input. Defaults to output_path.with_suffix('.report.json').",
    )
    parser.add_argument("--no-normalize-text", action="store_true", help="Do not normalize punctuation or spacing")
    parser.add_argument("--page-numbers", action="store_true", help="Insert conservative Word PAGE field page numbers when safe")
    parser.add_argument("--format-tables", action="store_true", help="Preserve source tables and conservatively format cell internals")
    parser.add_argument("--generate-toc", action="store_true", help="Generate a Word TOC field when heading hierarchy is clear")
    parser.add_argument("--format-imprint", action="store_true", help="Detect and format existing imprint lines without creating new ones")
    parser.add_argument("--generic-formal-text", action="store_true", help="Format as 通用正式文本: preserve order and apply generic formal typography without official-document structure")
    parser.add_argument("--standard-text", action="store_true", help="Format 标准规范文本 documents with cover, toc, chapter, clause, and table rules")
    parser.add_argument("--standard-spec-text", action="store_true", help="Alias for --standard-text")
    parser.add_argument(
        "--space-mode",
        choices=("keep_en_boundary", "remove_all", "keep_all"),
        default="keep_en_boundary",
        help="How to handle spaces when normalizing text.",
    )
    args = parser.parse_args()
    if args.generic_formal_text:
        args.doc_type = GENERIC_FORMAL_TEXT
    if args.standard_spec_text:
        args.standard_text = True

    profile = load_profile(args.profile)
    output_path = Path(args.output)
    report_path = _report_path(output_path, args.report)
    normalize = not args.no_normalize_text
    report_data = None
    footer_inspection = None
    input_path = None
    inline_table_count = 0
    chapter_recovery_method = None

    if args.create_skeleton:
        if not args.doc_type:
            raise SystemExit("--doc-type is required with --create-skeleton")
        title = args.title or f"关于××事项的{args.doc_type}"
        body = build_skeleton_body(args.doc_type)
        doc = build_document(title, args.recipient, body, args.issuer, args.date_text, profile, normalize=normalize, space_mode=args.space_mode)
    elif args.text_file:
        text = Path(args.text_file).read_text(encoding="utf-8")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        title = args.title or (lines[0] if lines else "")
        body = lines[1:] if title and lines and title == lines[0] else lines
        doc = build_document(title, args.recipient, body, args.issuer, args.date_text, profile, normalize=normalize, space_mode=args.space_mode)
    else:
        if not args.input:
            raise SystemExit("input .docx path is required unless --text-file is used")
        input_path = Path(args.input)
        if args.standard_text or (not args.generic_formal_text and looks_like_standard_text(input_path)):
            doc, standard_result = build_standard_text_document(input_path, profile)
            if report_path is not None:
                read_docx_snapshot, analyze_structure, diagnose_snapshot, build_format_plan, write_report_json = _engine_imports()
                snapshot = read_docx_snapshot(input_path)
                structure = analyze_structure(snapshot)
                report_data = {
                    "write_report_json": write_report_json,
                    "input_path": input_path,
                    "profile_id": profile.get("profile_id", args.profile),
                    "doc_type": STANDARD_SPEC_TEXT,
                    "structure": structure,
                    "issues": diagnose_snapshot(snapshot, structure),
                    "operations": [
                        FormatOperation(
                            kind="standard_text_format",
                            target="document",
                            params={
                                "action": standard_result.action,
                                "toc_count": standard_result.toc_count,
                                "table_count": standard_result.table_count,
                                "merged_cover_label": standard_result.merged_cover_label,
                            },
                            reason="apply standard-specification cover, toc hierarchy, body heading, and serial-number table rules",
                        )
                    ],
                    "FormatOperation": FormatOperation,
                }
            args.page_numbers = False
            args.format_tables = False
            args.generate_toc = False
            args.format_imprint = False
        else:
            read_docx_snapshot, analyze_structure, diagnose_snapshot, build_format_plan, write_report_json = _engine_imports()
            snapshot = repair_glued_snapshot(read_docx_snapshot(input_path))
            if args.page_numbers:
                footer_inspection = inspect_footers(input_path)
            source_lines = [paragraph.text.strip() for paragraph in snapshot.non_empty_paragraphs]
            if not args.doc_type:
                classification = classify_lines_for_type(source_lines)
                top = classification["top"]
                print(f"detected_doc_type={top.get('doc_type', '未知')} confidence={top.get('confidence', 0)}")
                if classification["ask_user"] and not args.assume_detected_type:
                    print(classification["question"])
                    return 2
                args.doc_type = top.get("doc_type")

            structure = analyze_structure(snapshot)
            issues = diagnose_snapshot(snapshot, structure)
            plan = build_format_plan(
                snapshot,
                structure,
                profile_id=profile.get("profile_id", args.profile),
                doc_type=args.doc_type or "未知",
                normalize_text=normalize,
            )
            doc, inline_table_count, chapter_recovery_method = build_document_from_source(
                input_path,
                args.title,
                args.recipient,
                args.issuer,
                args.date_text,
                profile,
                normalize=normalize,
                space_mode=args.space_mode,
                generic_formal_text=args.doc_type == GENERIC_FORMAL_TEXT,
            )
            if report_path is not None:
                report_data = {
                    "write_report_json": write_report_json,
                    "input_path": input_path,
                    "profile_id": profile.get("profile_id", args.profile),
                    "doc_type": args.doc_type or "未知",
                    "structure": structure,
                    "issues": issues,
                    "operations": list(plan.operations),
                    "FormatOperation": FormatOperation,
                }
                if chapter_recovery_method:
                    report_data["operations"].append(
                        FormatOperation(
                            kind="chapter_recovery",
                            target="paragraph",
                            params={
                                "method": chapter_recovery_method,
                                "chapters": sum(
                                    1
                                    for paragraph in doc.paragraphs
                                    if re.match(r"^[一二三四五六七八九十]+、", paragraph.text.strip())
                                ),
                            },
                            reason="split glued single-paragraph draft into deterministic chapter blocks",
                        )
                    )

    if args.format_tables and input_path is not None and inline_table_count == 0:
        table_result = append_and_format_source_tables(doc, input_path, profile)
        if report_data is not None:
            report_data["operations"].append(
                report_data["FormatOperation"](
                    kind="table_format",
                    target="tables",
                    params={
                        "table_count": table_result.table_count,
                        "skipped_count": table_result.skipped_count,
                        "action": table_result.action,
                    },
                    reason="preserve source tables and format cell internals conservatively",
                )
            )
    elif args.format_tables and input_path is not None and inline_table_count:
        if report_data is not None:
            report_data["operations"].append(
                report_data["FormatOperation"](
                    kind="table_format",
                    target="tables",
                    params={
                        "table_count": inline_table_count,
                        "skipped_count": 0,
                        "action": "formatted_inline",
                    },
                    reason="preserve source tables in original document order and format cell internals conservatively",
                )
            )

    if args.generate_toc:
        toc_result = generate_toc_if_clear(doc)
        if report_data is not None:
            report_data["operations"].append(
                report_data["FormatOperation"](
                    kind="toc_generation",
                    target="document",
                    params={
                        "heading_count": toc_result.heading_count,
                        "max_level": toc_result.max_level,
                        "action": toc_result.action,
                    },
                    reason="generate Word TOC field only when heading hierarchy is clear",
                )
            )

    if args.format_imprint:
        imprint_result = format_existing_imprint(doc, profile)
        if report_data is not None:
            report_data["operations"].append(
                report_data["FormatOperation"](
                    kind="imprint_format",
                    target="imprint",
                    params={
                        "imprint_count": imprint_result.imprint_count,
                        "paragraph_indices": list(imprint_result.paragraph_indices),
                        "action": imprint_result.action,
                    },
                    reason="format existing imprint lines without creating new imprint text",
                )
            )

    if args.page_numbers:
        page_number_result = apply_page_numbers(doc, footer_inspection)
        if report_data is not None:
            report_data["operations"].append(
                report_data["FormatOperation"](
                    kind="page_number",
                    target="footer",
                    params={
                        "existing_page_number": page_number_result.existing_page_number,
                        "existing_non_page_footer": page_number_result.existing_non_page_footer,
                        "action": page_number_result.action,
                        "preserved_footer_texts": list(page_number_result.preserved_footer_texts),
                    },
                    reason="insert PAGE field unless footer contains non-page content",
                )
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if report_path is not None and report_data is not None:
        report_data["write_report_json"](
            report_path,
            input_path=report_data["input_path"],
            output_path=output_path,
            profile_id=report_data["profile_id"],
            doc_type=report_data["doc_type"],
            structure=report_data["structure"],
            issues=report_data["issues"],
            operations=report_data["operations"],
        )
        print(f"report={report_path}")
    print(f"saved={output_path}")
    print(f"profile={profile.get('profile_id', args.profile)}")
    print(f"text_normalization={'off' if args.no_normalize_text else args.space_mode}")
    print(smoke_check(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
